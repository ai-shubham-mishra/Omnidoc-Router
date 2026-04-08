"""
Text Extraction Component for Omnidoc Router.
Extracts text content from various file types for embedding in Pinecone.

Supported formats:
- Text files: .txt, .md, .csv, .json, .xml, etc.
- PDFs: Uses PyPDF2 or pdfplumber
- Images: Placeholder for OCR (can integrate with Gemini Vision)
- Office docs: .docx, .xlsx (basic support)
"""
import os
import logging
from typing import Optional

logger = logging.getLogger(__name__)


class TextExtractor:
    """Extract text content from files for vector embedding."""
    
    @staticmethod
    def extract_text_from_bytes(content: bytes, mime_type: str, max_chars: int = 100000) -> Optional[str]:
        """Extract text directly from in-memory file bytes (no disk I/O)."""
        try:
            if mime_type.startswith("text/") or mime_type in [
                "application/json", "application/xml", "application/csv",
            ]:
                return content.decode("utf-8", errors="ignore")[:max_chars].strip()

            elif mime_type == "application/pdf":
                return TextExtractor._extract_pdf_bytes(content, max_chars)

            elif mime_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword",
            ]:
                return TextExtractor._extract_docx_bytes(content, max_chars)

            elif mime_type.startswith("image/"):
                logger.debug(f"Image file - OCR not implemented yet")
                return None

            else:
                logger.debug(f"Unsupported MIME type for text extraction: {mime_type}")
                return None
        except Exception as e:
            logger.error(f"Text extraction from bytes failed: {e}")
            return None

    @staticmethod
    def extract_text(file_path: str, mime_type: str, max_chars: int = 100000) -> Optional[str]:
        """Extract text from a local file path (fallback for local storage)."""
        try:
            if mime_type.startswith("text/") or mime_type in [
                "application/json", "application/xml", "application/csv",
            ]:
                return TextExtractor._extract_text_file(file_path, max_chars)
            elif mime_type == "application/pdf":
                return TextExtractor._extract_pdf(file_path, max_chars)
            elif mime_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword",
            ]:
                return TextExtractor._extract_docx(file_path, max_chars)
            elif mime_type.startswith("image/"):
                logger.debug(f"Image file detected - OCR not implemented yet")
                return None
            else:
                logger.debug(f"Unsupported MIME type: {mime_type}")
                return None
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {e}")
            return None

    @staticmethod
    def _extract_pdf_bytes(content: bytes, max_chars: int) -> Optional[str]:
        """Extract text from PDF bytes using PyPDF2."""
        try:
            import PyPDF2
            import io
            text_parts = []
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            for page in reader.pages:
                text_parts.append(page.extract_text())
                if sum(len(t) for t in text_parts) >= max_chars:
                    break
            return " ".join(text_parts)[:max_chars].strip()
        except ImportError:
            logger.warning("PyPDF2 not installed - PDF extraction disabled")
            return None
        except Exception as e:
            logger.error(f"PDF bytes extraction failed: {e}")
            return None

    @staticmethod
    def _extract_docx_bytes(content: bytes, max_chars: int) -> Optional[str]:
        """Extract text from DOCX bytes using python-docx."""
        try:
            from docx import Document
            import io
            doc = Document(io.BytesIO(content))
            text_parts = [p.text for p in doc.paragraphs]
            return " ".join(text_parts)[:max_chars].strip()
        except ImportError:
            logger.warning("python-docx not installed - DOCX extraction disabled")
            return None
        except Exception as e:
            logger.error(f"DOCX bytes extraction failed: {e}")
            return None

    @staticmethod
    def _extract_text_file(file_path: str, max_chars: int) -> Optional[str]:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read(max_chars).strip()
        except Exception as e:
            logger.error(f"Failed to read text file: {e}")
            return None

    @staticmethod
    def _extract_pdf(file_path: str, max_chars: int) -> Optional[str]:
        try:
            import PyPDF2
            text_parts = []
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text_parts.append(page.extract_text())
                    if sum(len(t) for t in text_parts) >= max_chars:
                        break
            return " ".join(text_parts)[:max_chars].strip()
        except ImportError:
            logger.warning("PyPDF2 not installed")
            return None
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return None

    @staticmethod
    def _extract_docx(file_path: str, max_chars: int) -> Optional[str]:
        try:
            from docx import Document
            doc = Document(file_path)
            text_parts = [p.text for p in doc.paragraphs]
            return " ".join(text_parts)[:max_chars].strip()
        except ImportError:
            logger.warning("python-docx not installed")
            return None
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return None


# Global instance
text_extractor = TextExtractor()
